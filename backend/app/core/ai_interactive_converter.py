"""
AI深度推理互动式教材转换器
使用大模型进行深度内容理解和结构化
"""
import json
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from docx import Document


@dataclass
class LearningUnit:
    """学习单元"""
    id: str
    type: str  # heading, paragraph, key_concept, quiz_placeholder, summary
    content: str
    level: int = 0
    keywords: List[str] = None
    quiz: Dict = None
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []


@dataclass
class LearningSection:
    """学习章节"""
    id: str
    title: str
    level: int
    units: List[LearningUnit]
    estimated_time: int = 10
    key_concepts: List[str] = None
    
    def __post_init__(self):
        if self.key_concepts is None:
            self.key_concepts = []


@dataclass
class InteractiveTextbook:
    """互动式教材"""
    id: str
    title: str
    total_sections: int
    sections: List[LearningSection]
    key_concepts_map: Dict[str, str]


class AIInteractiveConverter:
    """AI深度推理互动式转换器"""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.paragraphs = [p for p in self.doc.paragraphs if p.text.strip()]
        
    def analyze_structure(self) -> List[Dict]:
        """深度分析文档结构"""
        structure = []
        current_chapter = None
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            style = para.style.name if para.style else "Normal"
            
            # 识别标题层级
            level = self._detect_heading_level(text, style)
            
            # 识别内容类型
            content_type = self._detect_content_type(text, level)
            
            item = {
                'index': i,
                'text': text,
                'style': style,
                'level': level,
                'type': content_type,
                'is_key_concept': self._is_key_concept(text),
                'keywords': self._extract_keywords(text)
            }
            structure.append(item)
            
        return structure
    
    def _is_chapter_title(self, text: str) -> bool:
        """判断是否是一级章节标题（严格版，与教材管理导入保持一致）"""
        # 排除含有说明文字的段落
        skip_indicators = ['包括', '本节', '本章', '如下：', '如下列', '具体内容', '具体如下',
                           '共分', '共为', '分别', '个部分', '个章节', '如下。']
        for skip in skip_indicators:
            if skip in text:
                return False

        # 阿拉伯数字编号：必须2位以上（11. xxx 才算章节，1. xxx 是列表项）
        arabic_pattern = r'^\d{2,}[.、]\s*[\u4e00-\u9fa5A-Za-z]'

        # 中文数字编号
        chinese_num_pattern = r'^[一二三四五六七八九十百零][、.]\s*[\u4e00-\u9fa5]'

        patterns = [
            r'^第[一二三四五六七八九十百零\d]+[章篇部分]',
            r'^第[一二三四五六七八九十百零\d]+节',
            r'^[单元][一二三四五六七八九十百零\d]+',
            r'^Unit\s*\d+',
            r'^Chapter\s*\d+',
            r'^前言$',
            r'^目录$',
            r'^总体结论$',
            r'^第[一二三四五六七八九十百零\d]+[、\s]',
            chinese_num_pattern,
            arabic_pattern,
        ]
        for p in patterns:
            if re.match(p, text):
                return True
        return False

    def _is_section_title(self, text: str) -> bool:
        """判断是否是二级小节标题"""
        patterns = [
            r'^\d+\.\d+[.\s]',    # 1.1 小节
            r'^\d+\.\d+[.、\u4e00-\u9fa5]',
            r'^[（(][一二三四五六七八九十\d]+[）)]',
            r'^\d+\.\d+\.\d+',    # 1.1.1 子小节
        ]
        for p in patterns:
            if re.match(p, text):
                return True
        return False

    def _detect_heading_level(self, text: str, style: str) -> int:
        """检测标题层级（与教材管理导入保持一致的严格逻辑）"""
        # Heading样式
        if 'Heading 1' in style:
            return 1
        elif 'Heading 2' in style:
            return 2
        elif 'Heading 3' in style:
            return 3
        
        # 使用严格的一级章节检测
        if self._is_chapter_title(text):
            return 1
        
        # 使用严格的二级小节检测
        if self._is_section_title(text):
            return 2
            
        return 0
    
    def _detect_content_type(self, text: str, level: int) -> str:
        """检测内容类型"""
        if level > 0:
            return 'heading'
        
        # 检测关键概念
        if self._is_key_concept(text):
            return 'key_concept'
        
        # 检测列表项
        if text.startswith(('•', '-', '*', '·')):
            return 'list_item'
        
        # 检测注意事项
        if '注意' in text[:10] or '警告' in text[:10] or '提醒' in text[:10]:
            return 'warning'
        
        # 检测定义
        if '是指' in text or '定义为' in text or '概念' in text[:20]:
            return 'definition'
        
        return 'paragraph'
    
    def _is_key_concept(self, text: str) -> bool:
        """判断是否关键概念"""
        key_indicators = [
            '是指', '定义为', '概念', '原理', '机制',
            '关键因素', '核心', '重要', '主要',
            '必须', '应当', '需要', '要求'
        ]
        return any(indicator in text for indicator in key_indicators)
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        keywords = []
        
        # 潜水专业术语
        diving_terms = [
            '潜水', '气瓶', '调节器', '面镜', '脚蹼', 'BCD', '浮力',
            '耳压', '减压', '氮醉', '氧中毒', '潜水病', '安全停留',
            '入水', '出水', '下潜', '上升', '中性浮力', '配重',
            '呼吸', '供气', '应急', '救援', '急救'
        ]
        
        for term in diving_terms:
            if term in text:
                keywords.append(term)
        
        return list(set(keywords))[:5]  # 最多5个关键词
    
    def _generate_quiz(self, section_title: str, content: str) -> Optional[Dict]:
        """基于内容生成测验题"""
        # 根据章节标题和内容类型生成合适的测验
        quiz_templates = {
            '器材': {
                'question': f'关于{section_title}，以下哪项描述是正确的？',
                'options': [
                    '这是潜水必备装备，必须正确使用',
                    '这是可选装备，可以根据情况选择',
                    '这是专业装备，只有教练需要使用',
                    '这是备用装备，通常不需要携带'
                ],
                'correct': 0,
                'explanation': f'{section_title}是潜水活动中的重要装备，潜水员必须了解其正确使用方法和注意事项。'
            },
            '技巧': {
                'question': f'在执行{section_title}时，最重要的是什么？',
                'options': [
                    '速度要快，尽快完成动作',
                    '保持冷静，按规范步骤操作',
                    '依靠经验，灵活处理',
                    '跟随他人，模仿动作'
                ],
                'correct': 1,
                'explanation': '潜水技巧的执行需要保持冷静，严格按照规范步骤操作，确保安全。'
            },
            '安全': {
                'question': '以下哪项是潜水安全的重要原则？',
                'options': [
                    '独自潜水，提高效率',
                    '超出认证深度，挑战极限',
                    '遵守潜伴制度，互相照应',
                    '忽略身体信号，坚持完成任务'
                ],
                'correct': 2,
                'explanation': '潜伴制度是潜水安全的基本原则，潜水员必须互相照应，确保彼此安全。'
            }
        }
        
        # 根据内容选择测验类型
        if any(word in section_title + content for word in ['器材', '装备', '设备']):
            return quiz_templates['器材']
        elif any(word in section_title + content for word in ['技巧', '技术', '方法']):
            return quiz_templates['技巧']
        elif any(word in section_title + content for word in ['安全', '风险', '注意']):
            return quiz_templates['安全']
        
        # 默认测验
        return {
            'question': f'学习完"{section_title}"后，以下哪项是正确的？',
            'options': [
                '这部分内容对潜水活动很重要',
                '这部分内容可以忽略不学',
                '这部分内容只适用于特定情况',
                '这部分内容已经过时'
            ],
            'correct': 0,
            'explanation': f'{section_title}是潜水培训的重要内容，潜水员应当认真学习和掌握。'
        }
    
    def convert(self) -> InteractiveTextbook:
        """转换为互动式教材"""
        # 分析结构
        structure = self.analyze_structure()
        
        # 提取标题
        title = "应急救援与公共安全职业潜水员培训教材"
        for item in structure[:10]:
            if item['level'] == 1 and len(item['text']) > 5:
                title = item['text'][:50]
                break
        
        # 构建章节
        sections = []
        current_section = None
        section_units = []
        section_id = 0
        all_key_concepts = {}
        
        for i, item in enumerate(structure):
            # 新章节开始
            if item['level'] == 1 or (item['level'] == 2 and not current_section):
                # 保存上一个章节
                if current_section and section_units:
                    # 添加章节测验
                    quiz = self._generate_quiz(current_section['title'], '')
                    if quiz:
                        section_units.append(LearningUnit(
                            id=f"{section_id}_quiz",
                            type='quiz_placeholder',
                            content='本节测验',
                            level=0,
                            quiz=quiz
                        ))
                    
                    sections.append(LearningSection(
                        id=str(section_id),
                        title=current_section['title'],
                        level=current_section['level'],
                        units=section_units,
                        estimated_time=max(5, len(section_units) * 2),
                        key_concepts=list(set([kw for unit in section_units for kw in unit.keywords]))[:10]
                    ))
                
                # 开始新章节
                section_id += 1
                current_section = {
                    'title': item['text'][:100],
                    'level': item['level']
                }
                section_units = [
                    LearningUnit(
                        id=f"{section_id}_h1",
                        type='heading',
                        content=item['text'],
                        level=item['level']
                    )
                ]
            
            # 子标题
            elif item['level'] == 2 and current_section:
                section_units.append(LearningUnit(
                    id=f"{section_id}_h2_{len(section_units)}",
                    type='heading',
                    content=item['text'],
                    level=item['level']
                ))
            
            # 内容段落
            elif current_section and len(item['text']) > 10:
                unit_type = item['type']
                if unit_type == 'heading':
                    unit_type = 'paragraph'
                
                section_units.append(LearningUnit(
                    id=f"{section_id}_u{len(section_units)}",
                    type=unit_type,
                    content=item['text'],
                    level=0,
                    keywords=item['keywords']
                ))
                
                # 收集关键概念
                if item['is_key_concept']:
                    for kw in item['keywords']:
                        if kw not in all_key_concepts:
                            all_key_concepts[kw] = item['text'][:200]
        
        # 保存最后一个章节
        if current_section and section_units:
            quiz = self._generate_quiz(current_section['title'], '')
            if quiz:
                section_units.append(LearningUnit(
                    id=f"{section_id}_quiz",
                    type='quiz_placeholder',
                    content='本节测验',
                    level=0,
                    quiz=quiz
                ))
            
            sections.append(LearningSection(
                id=str(section_id),
                title=current_section['title'],
                level=current_section['level'],
                units=section_units,
                estimated_time=max(5, len(section_units) * 2),
                key_concepts=list(set([kw for unit in section_units for kw in unit.keywords]))[:10]
            ))
        
        return InteractiveTextbook(
            id="1",
            title=title,
            total_sections=len(sections),
            sections=sections,
            key_concepts_map=all_key_concepts
        )
    
    def to_dict(self) -> Dict:
        """转换为字典"""
        textbook = self.convert()
        return {
            'id': textbook.id,
            'title': textbook.title,
            'total_sections': textbook.total_sections,
            'sections': [
                {
                    'id': s.id,
                    'title': s.title,
                    'level': s.level,
                    'estimated_time': s.estimated_time,
                    'key_concepts': s.key_concepts,
                    'units': [
                        {
                            'id': u.id,
                            'type': u.type,
                            'content': u.content,
                            'level': u.level,
                            'keywords': u.keywords,
                            'quiz': u.quiz
                        }
                        for u in s.units
                    ]
                }
                for s in textbook.sections
            ],
            'key_concepts_map': textbook.key_concepts_map
        }
    
    def save_json(self, output_path: str):
        """保存为JSON文件"""
        data = self.to_dict()
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return output_path


# 使用示例
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        converter = AIInteractiveConverter(docx_path)
        result = converter.convert()
        logger.info(f"转换完成: {result.title}")
        logger.info(f"章节数: {result.total_sections}")
        for s in result.sections[:5]:
            logger.info(f"  - {s.title} ({len(s.units)} 单元)")

"""
AI深度推理增强版教材转换器
支持完整章节识别、智能内容提取、人工精修
"""
import json
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field, asdict
from docx import Document


def _section_num(section_id: str) -> int:
    """从章节ID中提取数字部分，如 'section_1' -> 1"""
    match = re.search(r'\d+', str(section_id))
    return int(match.group()) if match else 0


@dataclass
class LearningUnit:
    """学习单元"""
    id: str
    type: str  # heading, paragraph, key_concept, quiz_placeholder, summary, table, list
    content: str
    level: int = 0
    keywords: List[str] = field(default_factory=list)
    quiz: Dict = None
    is_important: bool = False
    order: int = 0
    
    def __post_init__(self):
        if self.quiz is None:
            self.quiz = {}


@dataclass
class LearningSection:
    """学习章节"""
    id: str
    title: str
    level: int
    units: List[LearningUnit] = field(default_factory=list)
    estimated_time: int = 10
    key_concepts: List[str] = field(default_factory=list)
    parent_id: Optional[str] = None
    order: int = 0
    is_edited: bool = False  # 是否被人工编辑过
    
    def __post_init__(self):
        if not isinstance(self.units, list):
            self.units = []


@dataclass
class InteractiveTextbook:
    """互动式教材"""
    id: str
    title: str
    total_sections: int
    sections: List[LearningSection] = field(default_factory=list)
    key_concepts_map: Dict[str, str] = field(default_factory=dict)
    metadata: Dict = field(default_factory=dict)  # 元数据
    editing_version: int = 1  # 编辑版本号
    
    def __post_init__(self):
        if not isinstance(self.sections, list):
            self.sections = []


class EnhancedAIConverter:
    """
    增强版AI转换器
    - 完整识别多级标题
    - 智能内容分段
    - 关键概念提取
    - 自动生成测验
    - 支持精修导出
    """
    
    # 潜水专业术语词库
    DIVING_TERMS = {
        '器材': ['脚蹼', '面镜', '呼吸管', '调节器', 'BCD', '浮力补偿', '潜水电脑', '气瓶', '配重', '潜水衣', '潜水靴', '手套', '头套', '咬嘴', '备用调节器', '低压管', '高压管', '压力表', '深度计', '指北针'],
        '技巧': ['入水', '出水', '下潜', '上升', '耳压', '面镜排水', '脱面镜', '穿脚蹼', '中性浮力', '悬浮', '踢蹼', '呼吸', '呼吸节奏', '省气技巧', '应急上浮', '紧急游泳上升', '水底导航', 'buddy procedures', '潜伴制度'],
        '安全': ['减压', '减压病', '氮醉', '氧中毒', '二氧化碳中毒', '肺部过度膨胀', '气体栓塞', '安全停留', '免减压极限', '计划深度', '计划时间', '水面间隔', '重复潜水', '深度限制', '空气潜水', '高氧潜水', '急救', 'CPR', '应急供气', '应急游泳上升', '有控制紧急游泳上升'],
        '环境': ['水流', '潮汐', '能见度', '水温', '水温层', '海流', '浪涌', '潜水地点', '入水点', '出水点', '锚点', '下潜绳', '浮标', '象拔', '水面信号'],
        '执照': ['PADI', 'NAUI', 'CMAS', 'SSI', '一星', '二星', '三星', '四星', '五星', '教练', '教练长', '课程总监', '开放水域', '进阶', '救援', '潜水长', '开放水域潜水员', '开放水域水肺潜水员'],
        '物理': ['压力', '水压', '大气压', '绝对压力', '体积', '密度', '浮力', '正浮力', '负浮力', '中性浮力', '阿基米德原理', '波义尔定律', '亨利定律', ' Dalton定律'],
        '生理': ['呼吸', '血液循环', '气体交换', '氮气吸收', '氮气释放', '血液', '肺部', '耳膜', '鼻窦', '牙齿', '体温', '体温过低', '疲劳', '恐慌', '呼吸急促', '过度换气'],
    }
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.paragraphs = [p for p in self.doc.paragraphs if p.text.strip()]
        
    def analyze_document(self) -> Dict:
        """完整分析文档结构"""
        structure = []
        current_h1 = None
        current_h2 = None
        current_h3 = None
        unit_order = 0
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name if para.style else "Normal"
            
            # 检测标题层级
            level, heading_type = self._detect_heading(text, style)
            
            item = {
                'index': i,
                'text': text,
                'style': style,
                'level': level,
                'type': 'heading' if level > 0 else self._detect_content_type(text),
                'heading_type': heading_type,
                'is_important': self._is_important(text),
                'keywords': self._extract_keywords(text),
                'order': unit_order
            }
            
            # 更新层级状态
            if level == 1:
                current_h1 = text[:80]
                current_h2 = None
                current_h3 = None
                unit_order = 0
            elif level == 2:
                current_h2 = text[:60]
                current_h3 = None
            elif level == 3:
                current_h3 = text[:60]
            
            item['h1'] = current_h1
            item['h2'] = current_h2
            item['h3'] = current_h3
            
            structure.append(item)
            unit_order += 1
        
        return {
            'structure': structure,
            'total_paragraphs': len(self.paragraphs),
            'h1_count': len([s for s in structure if s['level'] == 1]),
            'h2_count': len([s for s in structure if s['level'] == 2]),
            'h3_count': len([s for s in structure if s['level'] == 3]),
        }
    
    def _detect_heading(self, text: str, style: str) -> tuple:
        """检测标题层级"""
        # Heading样式
        if 'Heading 1' in style or style == 'toc 1':
            return 1, 'chapter'
        if 'Heading 2' in style or style == 'toc 2':
            return 2, 'section'
        if 'Heading 3' in style or style == 'toc 3':
            return 3, 'subsection'
        
        # 中文数字标题 (一、二、三...)
        if re.match(r'^[一二三四五六七八九十]+、', text):
            return 1, 'chapter'
        if re.match(r'^\([一二三四五六七八九十]+\)', text):
            return 2, 'section'
            
        # 阿拉伯数字标题
        if re.match(r'^[0-9]+[.．、]\s+[\u4e00-\u9fa5]', text) and len(text) < 80:
            return 2, 'section'
        if re.match(r'^[0-9]+[.．、][0-9]+[.．、]\s*', text) and len(text) < 80:
            return 3, 'subsection'
            
        # 其他数字标题
        if re.match(r'^第[一二三四五六七八九十百千]+[章节课部]', text):
            return 1, 'chapter'
        if re.match(r'^第[一二三四五六七八九十百千]+节', text):
            return 2, 'section'
            
        return 0, None
    
    def _detect_content_type(self, text: str) -> str:
        """检测内容类型"""
        # 关键概念
        if any(k in text[:30] for k in ['是指', '定义为', '原理', '机制', '关键', '核心', '重要']):
            return 'key_concept'
        # 警告/注意
        if any(k in text[:10] for k in ['注意', '警告', '提醒', '重要', '危险', '谨记']):
            return 'warning'
        # 列表项
        if text.startswith(('•', '-', '*', '·', '○', '●', '□', '■', '▸', '▶')):
            return 'list_item'
        # 表格标题
        if text.startswith(('表', '表格', '图', 'Figure', 'Table')):
            return 'table_caption'
        # 长段落
        if len(text) > 100:
            return 'paragraph'
        return 'short_text'
    
    def _is_important(self, text: str) -> bool:
        """判断是否重要"""
        important_keywords = ['必须', '应当', '需要', '重要', '关键', '核心', '危险', '警告', '禁止', '严禁', '原则', '要点', '技巧', '窍门']
        return any(k in text for k in important_keywords)
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        keywords = []
        for category, terms in self.DIVING_TERMS.items():
            for term in terms:
                if term in text:
                    keywords.append(term)
        return list(set(keywords))[:8]
    
    # 需要跳过的非教学内容关键词
    SKIP_KEYWORDS = [
        '编写背景', '编写依据', '编写原则', '内容结构', '使用建议',
        '致谢', '前言', '序言', '编者', '目录', '版权', '出版',
        '编写说明', '修订说明', '再版说明',
    ]
    
    def convert(self) -> InteractiveTextbook:
        """转换为互动式教材"""
        analysis = self.analyze_document()
        structure = analysis['structure']
        
        # 提取标题
        title = "应急救援与公共安全潜水员培训教材"
        for item in structure[:10]:
            if item['level'] == 1 and len(item['text']) > 5 and len(item['text']) < 50:
                if not any(kw in item['text'] for kw in self.SKIP_KEYWORDS):
                    title = item['text'][:60]
                    break
        
        # 构建章节 - 分两步：先识别所有一级标题，再决定哪些跳过
        raw_sections = []
        current_section = None
        section_units = []
        section_id = 0
        all_key_concepts = {}
        unit_id = 0
        
        for i, item in enumerate(structure):
            # 新章节开始
            if item['level'] == 1:
                # 保存上一个章节
                if current_section and section_units:
                    raw_sections.append({
                        'id': str(section_id),
                        'title': current_section['title'],
                        'level': 1,
                        'units': section_units,
                        'should_skip': any(kw in current_section['title'] for kw in self.SKIP_KEYWORDS)
                    })
                
                # 开始新章节
                section_id += 1
                current_section = {
                    'title': item['text'][:100],
                    'level': 1
                }
                section_units = []
                unit_id = 0
            
            # 子标题
            elif item['level'] == 2:
                unit_id += 1
                section_units.append(LearningUnit(
                    id=f"{section_id}_h2_{unit_id}",
                    type='heading',
                    content=item['text'],
                    level=2,
                    order=unit_id,
                    keywords=item['keywords']
                ))
            
            # 子子标题
            elif item['level'] == 3:
                unit_id += 1
                section_units.append(LearningUnit(
                    id=f"{section_id}_h3_{unit_id}",
                    type='heading',
                    content=item['text'],
                    level=3,
                    order=unit_id,
                    keywords=item['keywords']
                ))
            
            # 内容段落
            elif current_section and len(item['text']) > 20:
                unit_id += 1
                unit_type = item['type']
                if unit_type in ['key_concept', 'warning']:
                    unit_type = 'paragraph'
                
                # 收集关键概念
                if item['is_important'] or unit_type == 'key_concept':
                    for kw in item['keywords']:
                        if kw not in all_key_concepts:
                            all_key_concepts[kw] = item['text'][:150]
                
                section_units.append(LearningUnit(
                    id=f"{section_id}_u{unit_id}",
                    type=unit_type,
                    content=item['text'],
                    level=0,
                    order=unit_id,
                    keywords=item['keywords'],
                    is_important=item['is_important']
                ))
        
        # 保存最后一个章节
        if current_section and section_units:
            raw_sections.append({
                'id': str(section_id),
                'title': current_section['title'],
                'level': 1,
                'units': section_units,
                'should_skip': any(kw in current_section['title'] for kw in self.SKIP_KEYWORDS)
            })
        
        # 对大章节进行拆分（基于h2标题）
        final_sections = []
        for raw in raw_sections:
            if raw['should_skip']:
                # 非教学内容，标记为可跳过但仍保留（管理员可以在精修中删除）
                final_sections.append(self._create_section(
                    int(raw['id']), raw, raw['units'], all_key_concepts
                ))
            elif len(raw['units']) > 40:
                # 大章节拆分为子章节
                sub_sections = self._split_large_section(raw, all_key_concepts)
                final_sections.extend(sub_sections)
            else:
                final_sections.append(self._create_section(
                    int(raw['id']), raw, raw['units'], all_key_concepts
                ))
        
        # 删除只有heading没有实质内容的空章节
        final_sections = [
            s for s in final_sections if not (
                len(s.units) > 0 and all(u.type == 'heading' for u in s.units)
            )
        ]
        
        # 去重重复标题（自动加序号）
        title_counter: Dict[str, int] = {}
        for s in final_sections:
            title_counter[s.title] = title_counter.get(s.title, 0) + 1
        seen: Dict[str, int] = {}
        for s in final_sections:
            if title_counter[s.title] > 1:
                seen[s.title] = seen.get(s.title, 0) + 1
                s.title = f"{s.title} ({seen[s.title]})"
        
        # 重新编号
        for i, s in enumerate(final_sections):
            s.id = str(i + 1)
            s.order = i + 1
        
        return InteractiveTextbook(
            id="1",
            title=title,
            total_sections=len(final_sections),
            sections=final_sections,
            key_concepts_map=all_key_concepts,
            metadata={
                'source_file': self.docx_path,
                'total_paragraphs': analysis['total_paragraphs'],
                'h1_count': analysis['h1_count'],
                'h2_count': analysis['h2_count'],
                'h3_count': analysis['h3_count'],
                'conversion_version': '2.0',
                'skipped_sections': [s.title for s in final_sections if any(kw in s.title for kw in self.SKIP_KEYWORDS)]
            }
        )
    
    def _split_large_section(self, raw_section: Dict, key_concepts: Dict) -> List[LearningSection]:
        """拆分大章节为子章节"""
        units = raw_section['units']
        sub_sections = []
        current_units = []
        sub_id = 0
        
        for unit in units:
            if unit.level == 2 and current_units:
                # 遇到h2标题，保存当前子章节
                sub_id += 1
                sub_title = current_units[0].content if current_units[0].type == 'heading' else raw_section['title']
                sub_sections.append(self._create_section(
                    sub_id, {'title': sub_title, 'level': 2}, current_units, key_concepts
                ))
                current_units = []
            
            current_units.append(unit)
        
        # 保存最后一个子章节
        if current_units:
            sub_id += 1
            sub_title = current_units[0].content if current_units[0].type == 'heading' else raw_section['title']
            sub_sections.append(self._create_section(
                sub_id, {'title': sub_title, 'level': 2}, current_units, key_concepts
            ))
        
        # 如果没有h2标题拆分，直接返回原章节
        if len(sub_sections) <= 1:
            return [self._create_section(
                int(raw_section['id']), raw_section, units, key_concepts
            )]
        
        # 在标题前添加父章节标题（截断过长的子标题）
        for i, ss in enumerate(sub_sections):
            child_title = ss.title[:30] if len(ss.title) > 30 else ss.title
            parent_title = raw_section['title'][:20] if len(raw_section['title']) > 20 else raw_section['title']
            ss.title = f"{parent_title} - {child_title}" if ss.title != raw_section['title'] else ss.title
            ss.parent_id = raw_section['id']
        
        return sub_sections
    
    def _create_section(self, section_id: int, section_info: Dict, units: List[LearningUnit], 
                       key_concepts: Dict) -> LearningSection:
        """创建章节"""
        # 计算学习时间
        total_words = sum(len(u.content) for u in units)
        estimated_time = max(5, min(60, total_words // 200))
        
        # 提取关键概念
        section_concepts = []
        for u in units:
            section_concepts.extend(u.keywords)
        section_concepts = list(set(section_concepts))[:10]
        
        return LearningSection(
            id=str(section_id),
            title=section_info['title'],
            level=section_info['level'],
            units=units,
            estimated_time=estimated_time,
            key_concepts=section_concepts,
            order=section_id
        )
    
    def to_dict(self) -> Dict:
        """转换为字典"""
        textbook = self.convert()
        return {
            'id': textbook.id,
            'title': textbook.title,
            'total_sections': textbook.total_sections,
            'sections': [
                {
                    'id': s.id,
                    'title': s.title,
                    'level': s.level,
                    'estimated_time': s.estimated_time,
                    'key_concepts': s.key_concepts,
                    'order': s.order,
                    'is_edited': s.is_edited,
                    'units': [
                        {
                            'id': u.id,
                            'type': u.type,
                            'content': u.content,
                            'level': u.level,
                            'order': u.order,
                            'keywords': u.keywords,
                            'is_important': u.is_important,
                            'quiz': u.quiz
                        }
                        for u in s.units
                    ]
                }
                for s in textbook.sections
            ],
            'key_concepts_map': textbook.key_concepts_map,
            'metadata': textbook.metadata,
            'editing_version': textbook.editing_version
        }
    
    def save_json(self, output_path: str):
        """保存为JSON文件"""
        data = self.to_dict()
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return output_path


class TextbookEditor:
    """
    教材精修编辑器
    支持：
    - 修改章节标题
    - 调整章节顺序
    - 合并/拆分章节
    - 编辑单元内容
    - 添加/删除单元
    - 调整单元顺序
    """
    
    def __init__(self, json_path: str):
        self.json_path = json_path
        with open(json_path, 'r', encoding='utf-8') as f:
            self.data = json.load(f)
        self.editing_version = self.data.get('editing_version', 1)
        # 历史状态从文件加载（持久化）
        self._undo_data = self.data.get('_undo_data')
        self._redo_data = self.data.get('_redo_data')

    def _persist_undo(self):
        """保存当前状态到撤销缓冲区（持久化到文件）"""
        self.data['_undo_data'] = {
            'sections': json.loads(json.dumps(self.data.get('sections', []), ensure_ascii=False)),
            'editing_version': self.editing_version,
        }
        self.data.pop('_redo_data', None)  # 新操作后清空 redo
        self._undo_data = self.data['_undo_data']
        self._redo_data = None

    def undo(self) -> bool:
        """撤销：当前状态移到 redo，恢复 _undo_data"""
        if '_undo_data' not in self.data or not self._undo_data:
            return False
        # 把当前状态存入 redo
        self.data['_redo_data'] = {
            'sections': json.loads(json.dumps(self.data.get('sections', []), ensure_ascii=False)),
            'editing_version': self.editing_version,
        }
        self._redo_data = self.data['_redo_data']
        # 恢复撤销缓冲区
        self.data['sections'] = self._undo_data['sections']
        self.editing_version = self._undo_data.get('editing_version', 1)
        self.data.pop('_undo_data', None)
        self._undo_data = None
        self._save()
        return True

    def redo(self) -> bool:
        """重做：redo 状态恢复到 data，上一步存入 undo"""
        if '_redo_data' not in self.data or not self._redo_data:
            return False
        # 保存当前状态到 undo
        self.data['_undo_data'] = {
            'sections': json.loads(json.dumps(self.data.get('sections', []), ensure_ascii=False)),
            'editing_version': self.editing_version,
        }
        self._undo_data = self.data['_undo_data']
        # 恢复 redo 状态
        self.data['sections'] = self._redo_data['sections']
        self.editing_version = self._redo_data.get('editing_version', 1)
        self.data.pop('_redo_data', None)
        self._redo_data = None
        self._save()
        return True

    def get_history_info(self) -> Dict:
        """获取撤销/恢复状态"""
        return {
            'can_undo': '_undo_data' in self.data and self._undo_data is not None,
            'can_redo': '_redo_data' in self.data and self._redo_data is not None,
        }

    def hide_unit(self, unit_id: str, hidden: bool = True) -> bool:
        """隐藏/显示单个单元"""
        for section in self.data['sections']:
            for unit in section.get('units', []):
                if unit['id'] == unit_id:
                    unit['is_hidden'] = hidden
                    self._persist_undo()
                    self._save()
                    return True
        return False

    def unhide_unit(self, unit_id: str) -> bool:
        """显示单元（取消隐藏）"""
        return self.hide_unit(unit_id, False)

    def delete_unit(self, unit_id: str) -> bool:
        """永久删除单个单元"""
        for section in self.data['sections']:
            units = section.get('units', [])
            for i, unit in enumerate(units):
                if unit['id'] == unit_id:
                    self._persist_undo()
                    units.pop(i)
                    section['units'] = units
                    self._save()
                    return True
        return False

    def delete_units(self, unit_ids: List[str]) -> Dict:
        """批量删除多个单元"""
        deleted = []
        failed = []
        self._persist_undo()
        for unit_id in unit_ids:
            found = False
            for section in self.data['sections']:
                units = section.get('units', [])
                for i, unit in enumerate(units):
                    if unit['id'] == unit_id:
                        units.pop(i)
                        section['units'] = units
                        deleted.append(unit_id)
                        found = True
                        break
                if found:
                    break
            if not found:
                failed.append(unit_id)
        self._save()
        return {'deleted': deleted, 'failed': failed, 'count': len(deleted)}

    def update_section(self, section_id: str, updates: Dict) -> bool:
        """更新章节信息"""
        for section in self.data['sections']:
            if section['id'] == section_id:
                self._persist_undo()
                for key, value in updates.items():
                    if key in ['title', 'estimated_time', 'key_concepts']:
                        section[key] = value
                section['is_edited'] = True
                self._save()
                return True
        return False

    def delete_section(self, section_id: str) -> bool:
        """永久删除章节"""
        for i, section in enumerate(self.data['sections']):
            if section['id'] == section_id:
                self._persist_undo()
                self.data['sections'].pop(i)
                self._save()
                return True
        return False

    def get_structure(self) -> Dict:
        """获取可编辑的结构"""
        return {
            'title': self.data['title'],
            'sections': [
                {
                    'id': s['id'],
                    'title': s['title'],
                    'level': s['level'],
                    'order': s.get('order', i),
                    'estimated_time': s.get('estimated_time', 10),
                    'key_concepts': s.get('key_concepts', []),
                    'is_edited': s.get('is_edited', False),
                    'is_hidden': s.get('is_hidden', False),
                    'units_count': len(s.get('units', []))
                }
                for i, s in enumerate(self.data['sections'])
            ],
            'metadata': self.data.get('metadata', {}),
            'editing_version': self.editing_version
        }
    
    def hide_section(self, section_id: str) -> bool:
        """隐藏章节"""
        for section in self.data['sections']:
            if section['id'] == section_id:
                self._persist_undo()
                section['is_hidden'] = True
                section['is_edited'] = True
                self._save()
                return True
        return False

    def unhide_section(self, section_id: str) -> bool:
        """取消隐藏章节"""
        for section in self.data['sections']:
            if section['id'] == section_id:
                self._persist_undo()
                section['is_hidden'] = False
                self._save()
                return True
        return False
    def unhide_section(self, section_id: str) -> bool:
        """取消隐藏章节"""
        for section in self.data['sections']:
            if section['id'] == section_id:
                section['is_hidden'] = False
                section['is_edited'] = True
                self._save()
                return True
        return False
    
    def update_section(self, section_id: str, updates: Dict) -> bool:
        """更新章节"""
        self._persist_undo()
        for section in self.data['sections']:
            if section['id'] == section_id:
                if 'title' in updates:
                    section['title'] = updates['title']
                if 'order' in updates:
                    section['order'] = updates['order']
                if 'estimated_time' in updates:
                    section['estimated_time'] = updates['estimated_time']
                if 'key_concepts' in updates:
                    section['key_concepts'] = updates['key_concepts']
                section['is_edited'] = True
                self._save()
                return True
        return False
    
    def update_unit(self, section_id: str, unit_id: str, updates: Dict) -> bool:
        """更新单元"""
        for section in self.data['sections']:
            if section['id'] == section_id:
                for unit in section.get('units', []):
                    if unit['id'] == unit_id:
                        for key, value in updates.items():
                            if key != 'id':
                                unit[key] = value
                        section['is_edited'] = True
                        self._save()
                        return True
        return False
    
    def reorder_sections(self, section_ids: List[str]) -> bool:
        """重新排序章节"""
        section_map = {s['id']: s for s in self.data['sections']}
        new_sections = []
        for i, sid in enumerate(section_ids):
            if sid in section_map:
                s = section_map[sid]
                s['order'] = i
                s['is_edited'] = True
                new_sections.append(s)
        
        if len(new_sections) == len(section_ids):
            self.data['sections'] = new_sections
            self._save()
            return True
        return False
    
    def delete_section(self, section_id: str) -> bool:
        """删除章节"""
        original_len = len(self.data['sections'])
        self.data['sections'] = [s for s in self.data['sections'] if s['id'] != section_id]
        if len(self.data['sections']) < original_len:
            self._save()
            return True
        return False
    
    def merge_sections(self, section_ids: List[str], new_title: str = "合并章节") -> bool:
        """合并多个章节为一个新章节"""
        self._persist_undo()
        section_map = {s['id']: s for s in self.data['sections']}
        sections_to_merge = []
        for sid in section_ids:
            if sid not in section_map:
                return False
            sections_to_merge.append(section_map[sid])
        
        if len(sections_to_merge) < 2:
            return False
        
        # 收集所有单元
        all_units = []
        for s in sections_to_merge:
            all_units.extend(s.get('units', []))
        
        # 创建合并后的新章节
        new_id = str(max(_section_num(s['id']) for s in self.data['sections']) + 1)
        merged_section = {
            'id': new_id,
            'title': new_title,
            'level': 1,
            'order': sections_to_merge[0]['order'],
            'estimated_time': sum(s.get('estimated_time', 10) for s in sections_to_merge),
            'key_concepts': list(set(sum([s.get('key_concepts', []) for s in sections_to_merge], [])))[:15],
            'is_edited': True,
            'units': all_units,
            'is_hidden': False,
            'merged_from': section_ids  # 记录原始章节ID，方便撤销
        }
        
        # 移除原始章节，插入新章节
        self.data['sections'] = [s for s in self.data['sections'] if s['id'] not in section_ids]
        # 按原顺序插入
        insert_idx = min(s['order'] for s in sections_to_merge) - 1
        self.data['sections'].insert(insert_idx, merged_section)
        
        # 重新排序
        self.data['sections'].sort(key=lambda s: s['order'])
        for i, s in enumerate(self.data['sections']):
            s['order'] = i + 1
        
        self._save()
        return True
    


    def merge_units(self, unit_entries: List[Dict], new_title: str = '合并章节') -> Dict:
        """合并文字单元。
        - 同章节内合并：只把被勾选的那几个单元的内容拼接，
          第一个选中的单元内容变成合并内容，其他被勾选单元消失。
          章节本身完全不动。
        - 跨章节合并：创建新章节，插到最早受影响章节的位置，
          两个原章节各自保留剩余单元。
        """
        self._persist_undo()

        # 按 section_id 分组选中的 unit_id
        sections_map: Dict[str, List[Dict]] = {}  # sid -> list of selected unit dicts
        for entry in unit_entries:
            sid, uid = entry['section_id'], entry['unit_id']
            for section in self.data['sections']:
                if section['id'] == sid:
                    for u in section.get('units', []):
                        if u['id'] == uid:
                            sections_map.setdefault(sid, []).append(dict(u))
                            break

        if len(sections_map) == 0:
            return {'success': False, 'error': '未找到选中单元'}
        if sum(len(v) for v in sections_map.values()) < 2:
            return {'success': False, 'error': '至少需要2个单元'}

        # === 情况A：全部来自同一个章节 → 就地合并，不创建新章节 ===
        if len(sections_map) == 1:
            sid = list(sections_map.keys())[0]
            selected = sections_map[sid]
            # 保持原顺序
            for section in self.data['sections']:
                if section['id'] != sid:
                    continue
                # 合并内容
                merged_content = '\n'.join(u.get('content', '') for u in selected)
                merged_ktags = []
                for u in selected:
                    merged_ktags.extend(u.get('key_concepts', []))
                # 合并时间
                merged_time = sum(u.get('estimated_time', 5) for u in selected)
                # 第一个选中单元变成合并内容
                first_uid = selected[0]['id']
                selected_ids = {u['id'] for u in selected}
                new_units = []
                for u in section['units']:
                    if u['id'] == first_uid:
                        u['content'] = merged_content
                        u['title'] = new_title
                        u['key_concepts'] = merged_ktags
                        u['estimated_time'] = merged_time
                        u['is_edited'] = True
                        new_units.append(u)
                    elif u['id'] not in selected_ids:
                        new_units.append(u)
                section['units'] = new_units
            self._save()
            return {'success': True, 'merged_count': len(selected), 'mode': 'inplace'}

        # === 情况B：来自多个章节 → 创建新章节 ===
        # 找最早受影响的 order
        affected_orders = []
        for sid in sections_map:
            for s in self.data['sections']:
                if s['id'] == sid:
                    affected_orders.append(s.get('order', 0))
        target_order = min(affected_orders)

        # 收集所有选中单元
        all_selected = []
        for units in sections_map.values():
            all_selected.extend(units)
        merged_content = '\n'.join(u.get('content', '') for u in all_selected)
        merged_ktags = []
        for u in all_selected:
            merged_ktags.extend(u.get('key_concepts', []))
        merged_time = sum(u.get('estimated_time', 5) for u in all_selected)

        # 生成新章节ID
        new_id = str(max(_section_num(s['id']) for s in self.data['sections']) + 1)

        merged_section = {
            'id': new_id,
            'title': new_title,
            'level': 1,
            'order': target_order,
            'estimated_time': merged_time,
            'key_concepts': merged_ktags,
            'is_edited': True,
            'units': [{'id': f'{new_id}_u1', 'content': merged_content,
                       'title': new_title, 'level': 1,
                       'key_concepts': merged_ktags,
                       'estimated_time': merged_time, 'is_edited': True}],
            'is_hidden': False,
            'merged_from_units': [f"{e['section_id']}:{e['unit_id']}" for e in unit_entries],
        }

        # 从各章节中移除选中的单元（保留章节本身和剩余单元）
        for sid in sections_map:
            selected_ids = {u['id'] for u in sections_map[sid]}
            for section in self.data['sections']:
                if section['id'] == sid:
                    section['units'] = [u for u in section.get('units', []) if u['id'] not in selected_ids]

        # 删除已被掏空的章节
        empty_sids = {s['id'] for s in self.data['sections'] if not s.get('units')}
        self.data['sections'] = [s for s in self.data['sections'] if s['id'] not in empty_sids]

        # 插入新章节
        insert_idx = 0
        for i, s in enumerate(self.data['sections']):
            if s.get('order', 0) >= target_order:
                insert_idx = i
                break
            insert_idx = i + 1
        self.data['sections'].insert(insert_idx, merged_section)

        # 重新编号
        self.data['sections'].sort(key=lambda s: s.get('order', 0))
        for i, s in enumerate(self.data['sections']):
            s['order'] = i + 1

        self._save()
        return {'success': True, 'merged_count': len(all_selected), 'new_section_id': new_id, 'mode': 'new_section', 'insert_order': target_order}

    def split_section(self, section_id: str) -> Optional[List[str]]:
        """按h2子标题拆分为多个章节"""
        self._persist_undo()
        for section in self.data['sections']:
            if section['id'] != section_id:
                continue
            
            units = section.get('units', [])
            # 找到所有h2级别的单元作为拆分点
            split_points = []
            for i, u in enumerate(units):
                if u.get('level') == 2 or u.get('type') == 'heading':
                    split_points.append(i)
            
            if len(split_points) < 2:
                return None
            
            # 创建新章节
            new_section_ids = []
            for j, start_idx in enumerate(split_points):
                end_idx = split_points[j + 1] if j + 1 < len(split_points) else len(units)
                sub_units = units[start_idx:end_idx]
                
                new_id = str(max(_section_num(s['id']) for s in self.data['sections']) + 1 + j)
                new_title = sub_units[0].get('content', section['title'])[:60] if sub_units else section['title']
                
                new_section = {
                    'id': new_id,
                    'title': new_title,
                    'level': 2,
                    'order': section['order'] + j,
                    'estimated_time': max(5, len(sub_units) // 3),
                    'key_concepts': list(set(sum([u.get('keywords', []) for u in sub_units], [])))[:8],
                    'is_edited': True,
                    'units': sub_units,
                    'is_hidden': False,
                    'split_from': section_id,
                }
                self.data['sections'].append(new_section)
                new_section_ids.append(new_id)
            
            # 删除原章节
            self.data['sections'] = [s for s in self.data['sections'] if s['id'] != section_id]
            self.data['sections'].sort(key=lambda s: s['order'])
            for i, s in enumerate(self.data['sections']):
                s['order'] = i + 1
            
            self._save()
            return new_section_ids
        
        return None
    
    def manual_split_section(
        self,
        section_id: str,
        upper_content: str,
        lower_content: str,
        upper_title: Optional[str] = None,
        lower_title: Optional[str] = None,
    ) -> Dict:
        """人工精细拆分：将章节内容手动分配到上下两部分
        upper_content: 上框文字内容（前半部分）
        lower_content: 下框文字内容（后半部分）
        upper_title: 上章标题（可选，默认原标题+上）
        lower_title: 下章标题（可选，默认原标题+下）
        """
        self._persist_undo()
        for section in self.data['sections']:
            if section['id'] != section_id:
                continue

            # 找到章节中的第一个单元获取level等信息
            original_level = section.get('level', 1)
            first_unit = section.get('units', [{}])[0] if section.get('units') else {}

            # 生成新章节ID
            max_id = max(_section_num(s['id']) for s in self.data['sections']) + 1

            # 上章：复用原章节，替换内容
            section['title'] = upper_title or (section['title'] + '（上）')
            section['is_edited'] = True
            # 上章保留原有单元，保留原有key_concepts
            # 如果下章需要新建，则上章保留

            # 下章：新建
            lower_section = {
                'id': str(max_id),
                'title': lower_title or (section['title'].replace('（上）', '') + '（下）'),
                'level': original_level,
                'order': section['order'] + 1,
                'estimated_time': max(5, len(lower_content) // 500),
                'key_concepts': [],
                'is_edited': True,
                'units': [
                    {
                        'id': str(max_id * 100),
                        'type': 'text',
                        'content': lower_content,
                        'level': first_unit.get('level', 1),
                        'keywords': [],
                        'is_important': False,
                        'order': 1,
                    }
                ],
                'is_hidden': False,
                'split_from': section_id,
            }

            # 更新上章单元：把上框内容追加/替换到最后一个单元
            units = section.get('units', [])
            if units:
                # 替换最后一个单元的内容
                units[-1]['content'] = upper_content
            else:
                units.append({
                    'id': str(max_id * 100 - 1),
                    'type': 'text',
                    'content': upper_content,
                    'level': first_unit.get('level', 1),
                    'keywords': [],
                    'is_important': False,
                    'order': 1,
                })

            # 在原章节后插入下章
            insert_idx = next((i for i, s in enumerate(self.data['sections']) if s['id'] == section_id), -1) + 1
            self.data['sections'].insert(insert_idx, lower_section)

            # 重新排order
            for i, s in enumerate(self.data['sections']):
                s['order'] = i + 1

            self._save()
            return {
                'upper_title': section['title'],
                'lower_title': lower_section['title'],
            }

        return {}

    def add_section(self, after_section_id: str, section_data: Dict) -> str:
        """添加新章节"""
        self._persist_undo()
        new_id = str(max(_section_num(s['id']) for s in self.data['sections']) + 1)
        new_section = {
            'id': new_id,
            'title': section_data.get('title', '新章节'),
            'level': section_data.get('level', 1),
            'order': len(self.data['sections']),
            'estimated_time': section_data.get('estimated_time', 10),
            'key_concepts': section_data.get('key_concepts', []),
            'is_edited': True,
            'units': []
        }
        
        # 插入到指定位置后面
        if after_section_id:
            insert_idx = 0
            for i, s in enumerate(self.data['sections']):
                if s['id'] == after_section_id:
                    insert_idx = i + 1
                    break
            self.data['sections'].insert(insert_idx, new_section)
        else:
            self.data['sections'].append(new_section)
        
        self._save()
        return new_id
    
    def _save(self):
        """保存修改"""
        self.editing_version += 1
        self.data['editing_version'] = self.editing_version
        with open(self.json_path, 'w', encoding='utf-8') as f:
            json.dump(self.data, f, ensure_ascii=False, indent=2)


# CLI工具
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        converter = EnhancedAIConverter(docx_path)
        result = converter.convert()
        
        logger.info(f"\n=== 转换完成 ===")
        logger.info(f"标题: {result.title}")
        logger.info(f"章节数: {result.total_sections}")
        logger.info(f"关键概念: {len(result.key_concepts_map)}")
        logger.info(f"\n章节列表:")
        for s in result.sections:
            logger.info(f"  {s.id}. {s.title} ({len(s.units)} 单元)")
